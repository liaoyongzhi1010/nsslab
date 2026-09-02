from __future__ import annotations

import re
from html import escape
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.services.rich_text import RichTextRun, rich_text_blocks, rich_text_plain, sanitize_rich_text


PDF_FONT = "STSong-Light"
RUN_LABELS = {
    "knowledge_base": "知识库构建",
    "rag": "RAG A/B 对比",
    "agent": "Mini Crypto Agent",
}


def _plain(value: Any) -> str:
    text = str(value or "")
    text = re.sub(r"```[^\n]*\n?", "", text)
    text = text.replace("```", "").replace("**", "").replace("__", "").replace("`", "")
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    return text.strip()


def _pdf_text(value: Any) -> str:
    return escape(_plain(value)).replace("\n", "<br/>") or "—"


def _report_data(report: dict[str, Any], conclusion: str) -> dict[str, Any]:
    project = report["project"]
    details = report.get("details", {})
    runs = report.get("runs", [])
    rag_runs = [run for run in runs if run.get("type") == "rag"]
    agent_runs = [run for run in runs if run.get("type") == "agent"]
    stored_html = report.get("observation", {}).get("html", "")
    observation_html = sanitize_rich_text(conclusion) if conclusion.strip() else sanitize_rich_text(stored_html)
    return {
        "title": f"{project['name']} · 密码学实验报告",
        "project": project,
        "generated_at": report.get("generated_at", ""),
        "details": details,
        "runs": runs,
        "scores": report.get("scores", {}),
        "latest_rag": rag_runs[-1] if rag_runs else None,
        "latest_agent": agent_runs[-1] if agent_runs else None,
        "observation_html": observation_html,
        "observation_text": rich_text_plain(observation_html) or "尚未填写观察和感想。",
    }


def _pdf_styles() -> dict[str, ParagraphStyle]:
    try:
        pdfmetrics.getFont(PDF_FONT)
    except KeyError:
        pdfmetrics.registerFont(UnicodeCIDFont(PDF_FONT))
    sample = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "CryptoTitle", parent=sample["Title"], fontName=PDF_FONT, fontSize=22, leading=31,
            textColor=colors.HexColor("#12352D"), alignment=TA_CENTER, spaceAfter=8 * mm,
        ),
        "subtitle": ParagraphStyle(
            "CryptoSubtitle", parent=sample["Normal"], fontName=PDF_FONT, fontSize=9, leading=14,
            textColor=colors.HexColor("#5A6D68"), alignment=TA_CENTER, spaceAfter=7 * mm,
        ),
        "heading": ParagraphStyle(
            "CryptoHeading", parent=sample["Heading2"], fontName=PDF_FONT, fontSize=15, leading=22,
            textColor=colors.HexColor("#12352D"), spaceBefore=6 * mm, spaceAfter=3 * mm,
        ),
        "body": ParagraphStyle(
            "CryptoBody", parent=sample["BodyText"], fontName=PDF_FONT, fontSize=10.5, leading=18,
            textColor=colors.HexColor("#233B35"), spaceAfter=2.4 * mm,
        ),
        "small": ParagraphStyle(
            "CryptoSmall", parent=sample["BodyText"], fontName=PDF_FONT, fontSize=8.5, leading=13,
            textColor=colors.HexColor("#51655F"),
        ),
        "observation_heading": ParagraphStyle(
            "ObservationHeading", parent=sample["Heading3"], fontName=PDF_FONT, fontSize=12, leading=18,
            textColor=colors.HexColor("#215348"), spaceBefore=2.5 * mm, spaceAfter=1.5 * mm,
        ),
        "quote": ParagraphStyle(
            "ObservationQuote", parent=sample["BodyText"], fontName=PDF_FONT, fontSize=10.5, leading=18,
            textColor=colors.HexColor("#49665D"), leftIndent=6 * mm, borderColor=colors.HexColor("#9BC9BA"),
            borderWidth=1, borderPadding=5, spaceAfter=2.4 * mm,
        ),
    }


def _pdf_paragraph(value: Any, style: ParagraphStyle) -> Paragraph:
    return Paragraph(_pdf_text(value), style)


def _pdf_rich_text(value: str, styles: dict[str, ParagraphStyle]) -> list[Paragraph]:
    rendered: list[Paragraph] = []
    for block in rich_text_blocks(value):
        parts: list[str] = []
        for run in block.runs:
            text = escape(run.text).replace("\n", "<br/>")
            if run.bold:
                text = f"<b>{text}</b>"
            if run.italic:
                text = f"<i>{text}</i>"
            if run.underline:
                text = f"<u>{text}</u>"
            if run.strike:
                text = f"<strike>{text}</strike>"
            if run.href:
                text = f'<link href="{escape(run.href, quote=True)}" color="#315F8E">{text}</link>'
            parts.append(text)
        body = "".join(parts) or "—"
        style = styles["observation_heading"] if block.kind == "heading" else styles["quote"] if block.kind == "quote" else styles["body"]
        bullet = "•" if block.kind == "bullet" else f"{block.ordinal or 1}." if block.kind == "number" else None
        rendered.append(Paragraph(body, style, bulletText=bullet))
    return rendered or [_pdf_paragraph("尚未填写观察和感想。", styles["body"])]


def _pdf_table(rows: list[list[Any]], widths: list[float], styles: dict[str, ParagraphStyle], *, header: bool = False) -> Table:
    rendered = [[_pdf_paragraph(cell, styles["small"] if row_index else styles["body"]) for cell in row] for row_index, row in enumerate(rows)]
    table = Table(rendered, colWidths=widths, repeatRows=1 if header else 0, hAlign="LEFT")
    commands: list[tuple[Any, ...]] = [
        ("FONTNAME", (0, 0), (-1, -1), PDF_FONT),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C9D8D3")),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F0F8F5")),
    ]
    if header:
        commands += [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DDF3EB")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#12352D")),
        ]
    table.setStyle(TableStyle(commands))
    return table


def build_pdf(report: dict[str, Any], conclusion: str = "") -> bytes:
    data = _report_data(report, conclusion)
    styles = _pdf_styles()
    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer, pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title=data["title"], author="CryptoLLMLab AI 赋能密码学实验平台",
    )
    story: list[Any] = [
        _pdf_paragraph(data["title"], styles["title"]),
        _pdf_paragraph(f"生成时间：{data['generated_at']}　｜　由真实实验运行记录生成", styles["subtitle"]),
        _pdf_paragraph("一、项目概况", styles["heading"]),
    ]
    project = data["project"]
    stats = project.get("stats", {})
    story.append(_pdf_table([
        ["项目名称", project.get("name")],
        ["项目 ID", project.get("id")],
        ["当前阶段", f"Stage {project.get('current_stage', 1)}"],
        ["实验记录", f"{len(data['runs'])} 条"],
        ["知识库规模", f"{stats.get('documents', 0)} 份资料 / {stats.get('chunks', 0)} 个 Chunk"],
    ], [36 * mm, 138 * mm], styles))

    details = data["details"]
    kb = details.get("knowledge_base")
    story.append(_pdf_paragraph("二、实验一 · 密码学知识库", styles["heading"]))
    if kb:
        story.append(_pdf_table([
            ["Knowledge Base ID", kb.get("id")],
            ["Chunk 参数", f"Size {kb.get('chunk_size')} / Overlap {kb.get('overlap')}"],
            ["Embedding", f"{kb.get('embedding_model')} ({kb.get('dimension')}D)"],
            ["资料", "、".join(item.get("title", "") for item in details.get("documents", [])) or "—"],
        ], [36 * mm, 138 * mm], styles))
    else:
        story.append(_pdf_paragraph("尚未完成知识库构建。", styles["body"]))

    story.append(_pdf_paragraph("三、实验二 · Crypto-RAG", styles["heading"]))
    rag = details.get("rag_pipeline")
    if rag:
        story.append(_pdf_table([
            ["Pipeline ID", rag.get("id")],
            ["检索配置", f"Top-K {rag.get('top_k')} / Threshold {rag.get('threshold')} / Rerank Top-N {rag.get('rerank_top_n')}"],
            ["Prompt 模板", rag.get("prompt_template")],
        ], [36 * mm, 138 * mm], styles))
        latest_rag = data["latest_rag"]
        if latest_rag:
            output = latest_rag.get("output", {})
            story += [
                Spacer(1, 3 * mm),
                _pdf_paragraph(f"最近问题：{latest_rag.get('input', {}).get('query', '')}", styles["body"]),
                _pdf_paragraph("基础模型回答", styles["small"]),
                _pdf_paragraph(output.get("base", {}).get("answer", "—"), styles["body"]),
                _pdf_paragraph("RAG 回答与引用", styles["small"]),
                _pdf_paragraph(output.get("rag", {}).get("answer", "—"), styles["body"]),
            ]
    else:
        story.append(_pdf_paragraph("尚未运行 RAG A/B 对比。", styles["body"]))

    story.append(PageBreak())
    story.append(_pdf_paragraph("四、实验三 · Mini Crypto Agent", styles["heading"]))
    agent = details.get("agent")
    if agent:
        story.append(_pdf_table([
            ["Agent ID", agent.get("id")],
            ["启用 Skills", "、".join(agent.get("skills", []))],
            ["启用 Tools", "、".join(agent.get("tools", []))],
        ], [36 * mm, 138 * mm], styles))
        latest_agent = data["latest_agent"]
        if latest_agent:
            output = latest_agent.get("output", {})
            story += [
                Spacer(1, 3 * mm),
                _pdf_paragraph(f"最近任务：{latest_agent.get('input', {}).get('query', '')}", styles["body"]),
                _pdf_paragraph(f"选用 Skill：{output.get('selected_skill', {}).get('id', '—')}", styles["body"]),
                _pdf_paragraph("Agent 输出", styles["small"]),
                _pdf_paragraph(output.get("answer", "—"), styles["body"]),
            ]
    else:
        story.append(_pdf_paragraph("尚未运行 Mini Crypto Agent。", styles["body"]))

    story.append(_pdf_paragraph("五、能力画像", styles["heading"]))
    score_rows = [["能力维度", "得分"], *[[label, score] for label, score in data["scores"].items()]]
    story.append(_pdf_table(score_rows, [130 * mm, 44 * mm], styles, header=True))

    story.append(_pdf_paragraph("六、运行记录", styles["heading"]))
    history_rows: list[list[Any]] = [["实验", "输入 / 内容", "时间"]]
    for run in data["runs"]:
        input_data = run.get("input", {})
        summary = input_data.get("query") or f"{len(input_data.get('documents', []))} 份密码学资料"
        history_rows.append([RUN_LABELS.get(run.get("type", ""), run.get("type", "—")), summary, run.get("created_at", "")])
    if len(history_rows) == 1:
        history_rows.append(["—", "暂无运行记录", "—"])
    story.append(_pdf_table(history_rows, [38 * mm, 88 * mm, 48 * mm], styles, header=True))

    story.append(_pdf_paragraph("七、观察和感想", styles["heading"]))
    story.extend(_pdf_rich_text(data["observation_html"], styles))

    def footer(canvas: Any, doc: Any) -> None:
        canvas.saveState()
        canvas.setFont(PDF_FONT, 8)
        canvas.setFillColor(colors.HexColor("#70817C"))
        canvas.drawString(18 * mm, 10 * mm, "CryptoLLMLab · AI 赋能密码学实验平台")
        canvas.drawRightString(A4[0] - 18 * mm, 10 * mm, f"第 {doc.page} 页")
        canvas.restoreState()

    document.build(story, onFirstPage=footer, onLaterPages=footer)
    return buffer.getvalue()


def _set_docx_font(style: Any, size: float | None = None, *, color: str | None = None) -> None:
    style.font.name = "Aptos"
    if size:
        style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _docx_table(document: Document, rows: list[list[Any]], widths: list[float] | None = None, *, header: bool = False) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Shading Accent 1"
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = _plain(value) or "—"
            if widths:
                cell.width = Cm(widths[column_index])
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    if header and row_index == 0:
                        run.bold = True


def _docx_paragraph(document: Document, text: Any, *, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if bold_prefix:
        prefix = paragraph.add_run(bold_prefix)
        prefix.bold = True
    paragraph.add_run(_plain(text) or "—")


def _format_docx_run(run: Any, rich_run: RichTextRun) -> None:
    run.bold = rich_run.bold
    run.italic = rich_run.italic
    run.underline = rich_run.underline
    run.font.strike = rich_run.strike
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _docx_hyperlink(paragraph: Any, rich_run: RichTextRun) -> None:
    relationship_id = paragraph.part.relate_to(rich_run.href, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = paragraph.add_run(rich_run.text)
    _format_docx_run(run, rich_run)
    run.font.color.rgb = RGBColor.from_string("315F8E")
    run.underline = True
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)


def _docx_rich_text(document: Document, value: str) -> None:
    blocks = rich_text_blocks(value)
    if not blocks:
        _docx_paragraph(document, "尚未填写观察和感想。")
        return
    for block in blocks:
        if block.kind == "heading":
            paragraph = document.add_heading("", level=2)
        elif block.kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
        elif block.kind == "number":
            paragraph = document.add_paragraph(style="List Number")
        elif block.kind == "quote":
            paragraph = document.add_paragraph(style="Quote")
        else:
            paragraph = document.add_paragraph()
        for rich_run in block.runs:
            if rich_run.href:
                _docx_hyperlink(paragraph, rich_run)
            else:
                run = paragraph.add_run(rich_run.text)
                _format_docx_run(run, rich_run)


def build_docx(report: dict[str, Any], conclusion: str = "") -> bytes:
    data = _report_data(report, conclusion)
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    document.core_properties.title = data["title"]
    document.core_properties.author = "CryptoLLMLab AI 赋能密码学实验平台"
    for style_name, size, color in [("Normal", 10.5, "233B35"), ("Title", 22, "12352D"), ("Heading 1", 15, "12352D"), ("Heading 2", 12, "215348")]:
        _set_docx_font(document.styles[style_name], size, color=color)

    title = document.add_heading(data["title"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(f"生成时间：{data['generated_at']}\n由真实实验运行记录生成")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    project = data["project"]
    stats = project.get("stats", {})
    document.add_heading("一、项目概况", level=1)
    _docx_table(document, [
        ["项目名称", project.get("name")],
        ["项目 ID", project.get("id")],
        ["当前阶段", f"Stage {project.get('current_stage', 1)}"],
        ["实验记录", f"{len(data['runs'])} 条"],
        ["知识库规模", f"{stats.get('documents', 0)} 份资料 / {stats.get('chunks', 0)} 个 Chunk"],
    ], [4, 12.5])

    details = data["details"]
    kb = details.get("knowledge_base")
    document.add_heading("二、实验一 · 密码学知识库", level=1)
    if kb:
        _docx_table(document, [
            ["Knowledge Base ID", kb.get("id")],
            ["Chunk 参数", f"Size {kb.get('chunk_size')} / Overlap {kb.get('overlap')}"],
            ["Embedding", f"{kb.get('embedding_model')} ({kb.get('dimension')}D)"],
            ["资料", "、".join(item.get("title", "") for item in details.get("documents", [])) or "—"],
        ], [4, 12.5])
    else:
        _docx_paragraph(document, "尚未完成知识库构建。")

    document.add_heading("三、实验二 · Crypto-RAG", level=1)
    rag = details.get("rag_pipeline")
    if rag:
        _docx_table(document, [
            ["Pipeline ID", rag.get("id")],
            ["检索配置", f"Top-K {rag.get('top_k')} / Threshold {rag.get('threshold')} / Rerank Top-N {rag.get('rerank_top_n')}"],
            ["Prompt 模板", rag.get("prompt_template")],
        ], [4, 12.5])
        latest_rag = data["latest_rag"]
        if latest_rag:
            output = latest_rag.get("output", {})
            _docx_paragraph(document, latest_rag.get("input", {}).get("query", ""), bold_prefix="最近问题：")
            document.add_heading("基础模型回答", level=2)
            _docx_paragraph(document, output.get("base", {}).get("answer", "—"))
            document.add_heading("RAG 回答与引用", level=2)
            _docx_paragraph(document, output.get("rag", {}).get("answer", "—"))
    else:
        _docx_paragraph(document, "尚未运行 RAG A/B 对比。")

    document.add_heading("四、实验三 · Mini Crypto Agent", level=1)
    agent = details.get("agent")
    if agent:
        _docx_table(document, [
            ["Agent ID", agent.get("id")],
            ["启用 Skills", "、".join(agent.get("skills", []))],
            ["启用 Tools", "、".join(agent.get("tools", []))],
        ], [4, 12.5])
        latest_agent = data["latest_agent"]
        if latest_agent:
            output = latest_agent.get("output", {})
            _docx_paragraph(document, latest_agent.get("input", {}).get("query", ""), bold_prefix="最近任务：")
            _docx_paragraph(document, output.get("selected_skill", {}).get("id", "—"), bold_prefix="选用 Skill：")
            document.add_heading("Agent 输出", level=2)
            _docx_paragraph(document, output.get("answer", "—"))
    else:
        _docx_paragraph(document, "尚未运行 Mini Crypto Agent。")

    document.add_heading("五、能力画像", level=1)
    _docx_table(document, [["能力维度", "得分"], *[[label, score] for label, score in data["scores"].items()]], [12.5, 4], header=True)

    document.add_heading("六、运行记录", level=1)
    history_rows: list[list[Any]] = [["实验", "输入 / 内容", "时间"]]
    for run in data["runs"]:
        input_data = run.get("input", {})
        summary = input_data.get("query") or f"{len(input_data.get('documents', []))} 份密码学资料"
        history_rows.append([RUN_LABELS.get(run.get("type", ""), run.get("type", "—")), summary, run.get("created_at", "")])
    if len(history_rows) == 1:
        history_rows.append(["—", "暂无运行记录", "—"])
    _docx_table(document, history_rows, [3.5, 8.5, 4.5], header=True)

    document.add_heading("七、观察和感想", level=1)
    _docx_rich_text(document, data["observation_html"])

    footer = section.footer.paragraphs[0]
    footer.text = "CryptoLLMLab · AI 赋能密码学实验平台"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
