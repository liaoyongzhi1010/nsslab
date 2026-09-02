from fastapi.testclient import TestClient
from io import BytesIO
from docx import Document
from pypdf import PdfReader
from pypdf import PdfWriter
from pypdf.generic import DictionaryObject, NameObject, StreamObject
import pytest

from app.main import app, auth_service
from app.services.platform import platform_service


client = TestClient(app)


def setup_function():
    platform_service.reset()
    auth_service.ensure_bootstrap_users()
    client.cookies.clear()
    login = client.post(
        "/api/auth/login",
        json={"username": "admin", "password": "Admin-Test-Password-2026!"},
    )
    assert login.status_code == 200


def test_authentication_roles_logout_and_project_ownership():
    client.cookies.clear()
    assert client.get("/api/bootstrap").status_code == 401
    wrong = client.post(
        "/api/auth/login", json={"username": "student", "password": "definitely-wrong"}
    )
    assert wrong.status_code == 401

    student_login = client.post(
        "/api/auth/login",
        json={"username": "student", "password": "Student-Test-Password-2026!"},
    )
    assert student_login.status_code == 200
    assert student_login.json()["user"]["role"] == "student"
    assert "httponly" in student_login.headers["set-cookie"].lower()
    assert "samesite=strict" in student_login.headers["set-cookie"].lower()
    student_project = client.post("/api/projects", json={"name": "学生私有实验"})
    assert student_project.status_code == 201
    student_project_id = student_project.json()["id"]
    assert (
        client.put(
            "/api/skills/crypto_explain",
            json={
                "description": "学生不应修改全局技能",
                "steps": ["blocked"],
                "enabled": True,
            },
        ).status_code
        == 403
    )
    assert client.get("/api/admin/users").status_code == 403
    assert client.post("/api/auth/logout").status_code == 204
    assert client.get(f"/api/projects/{student_project_id}").status_code == 401

    admin_login = client.post(
        "/api/auth/login",
        json={"username": "admin", "password": "Admin-Test-Password-2026!"},
    )
    assert admin_login.status_code == 200
    assert admin_login.json()["user"]["role"] == "admin"
    assert client.get(f"/api/projects/{student_project_id}").status_code == 200
    created_user = client.post(
        "/api/admin/users",
        json={
            "username": "student2",
            "display_name": "学生二号",
            "password": "Student-Two-Password-2026!",
            "role": "student",
        },
    )
    assert created_user.status_code == 201

    client.post("/api/auth/logout")
    assert (
        client.post(
            "/api/auth/login",
            json={"username": "student2", "password": "Student-Two-Password-2026!"},
        ).status_code
        == 200
    )
    assert client.get(f"/api/projects/{student_project_id}").status_code == 403
    assert (
        client.patch(
            f"/api/projects/{student_project_id}", json={"ended": True}
        ).status_code
        == 403
    )


def test_public_registration_creates_only_student_and_starts_session():
    client.cookies.clear()
    registered = client.post(
        "/api/auth/register",
        json={
            "username": "new.student",
            "display_name": "新同学",
            "password": "New-Student-Password-2026!",
        },
    )
    assert registered.status_code == 201
    assert registered.json()["user"]["role"] == "student"
    assert registered.json()["user"]["display_name"] == "新同学"
    assert "httponly" in registered.headers["set-cookie"].lower()
    assert client.get("/api/auth/me").json()["user"]["username"] == "new.student"
    assert (
        client.post("/api/projects", json={"name": "注册用户项目"}).status_code == 201
    )

    duplicate = client.post(
        "/api/auth/register",
        json={
            "username": "new.student",
            "display_name": "重复账号",
            "password": "Another-Password-2026!",
        },
    )
    assert duplicate.status_code == 400
    injected_role = client.post(
        "/api/auth/register",
        json={
            "username": "fake-admin",
            "display_name": "越权账号",
            "password": "Fake-Admin-Password-2026!",
            "role": "admin",
        },
    )
    assert injected_role.status_code == 422
    assert auth_service.user_by_username("fake-admin") is None


def test_restart_moves_current_experiment_to_history_and_restore_is_recoverable():
    client.cookies.clear()
    assert (
        client.post(
            "/api/auth/login",
            json={"username": "student", "password": "Student-Test-Password-2026!"},
        ).status_code
        == 200
    )
    first = client.post("/api/projects", json={"name": "第一次 RAG 实验"}).json()
    built = client.post(
        "/api/kb/build",
        json={
            "project_id": first["id"],
            "document_ids": ["aes", "rsa"],
            "chunk_size": 256,
            "overlap": 32,
        },
    )
    assert built.status_code == 200

    second = client.post("/api/projects", json={"name": "第二次 Agent 实验"})
    assert second.status_code == 201
    current = client.get("/api/projects").json()
    assert [item["id"] for item in current] == [second.json()["id"]]
    history = client.get("/api/projects?include_ended=true").json()
    archived_first = next(item for item in history if item["id"] == first["id"])
    assert archived_first["is_ended"] is True
    assert archived_first["ended_at"] is not None
    assert archived_first["stats"]["knowledge_base"] is True
    assert archived_first["stats"]["runs"] == 1
    assert client.get(f"/api/reports/{first['id']}").status_code == 200
    assert (
        client.post(
            "/api/kb/build",
            json={
                "project_id": first["id"],
                "document_ids": ["aes"],
                "chunk_size": 256,
                "overlap": 32,
            },
        ).status_code
        == 400
    )

    renamed = client.patch(
        f"/api/projects/{first['id']}", json={"name": "第一次实验（已复盘）"}
    )
    assert renamed.status_code == 200
    assert renamed.json()["name"] == "第一次实验（已复盘）"
    restored = client.patch(f"/api/projects/{first['id']}", json={"ended": False})
    assert restored.status_code == 200
    assert restored.json()["is_ended"] is False
    after_restore = client.get("/api/projects").json()
    assert [item["id"] for item in after_restore] == [first["id"]]
    second_after_restore = client.get(f"/api/projects/{second.json()['id']}").json()
    assert second_after_restore["is_ended"] is True


def test_account_is_temporarily_locked_after_repeated_failures():
    created = client.post(
        "/api/admin/users",
        json={
            "username": "lock-test",
            "display_name": "锁定测试账号",
            "password": "Lock-Test-Password-2026!",
            "role": "student",
        },
    )
    assert created.status_code == 201
    client.post("/api/auth/logout")
    for _ in range(5):
        response = client.post(
            "/api/auth/login",
            json={"username": "lock-test", "password": "wrong-password"},
        )
        assert response.status_code == 401
    locked = client.post(
        "/api/auth/login",
        json={"username": "lock-test", "password": "Lock-Test-Password-2026!"},
    )
    assert locked.status_code == 401
    assert "登录失败次数过多" in locked.json()["detail"]


def create_project_and_kb():
    project = client.post("/api/projects", json={"name": "验收项目"}).json()
    response = client.post(
        "/api/kb/build",
        json={
            "project_id": project["id"],
            "document_ids": ["aes", "rsa", "tee", "he", "mpc"],
            "chunk_size": 512,
            "overlap": 64,
        },
    )
    assert response.status_code == 200
    return project["id"], response.json()


def test_bootstrap_and_build_knowledge_base():
    bootstrap = client.get("/api/bootstrap")
    assert bootstrap.status_code == 200
    body = bootstrap.json()
    assert len(body["documents"]) >= 25
    assert len(body["rag_benchmarks"]) == 10
    assert all("facts" not in benchmark for benchmark in body["rag_benchmarks"])
    assert (
        next(row for row in body["documents"] if row["id"] == "nist_fips203")[
            "source_type"
        ]
        == "NIST 正式标准"
    )
    assert (
        next(row for row in body["documents"] if row["id"] == "gmit_2023_revision")[
            "source_type"
        ]
        == "国家密码管理局官方公告"
    )
    assert next(row for row in body["documents"] if row["id"] == "kunpeng_secgear_dev")[
        "local_original"
    ].endswith(".html.txt")
    project_id, kb = create_project_and_kb()
    assert kb["document_count"] == 5
    assert kb["chunk_count"] >= 5
    assert kb["dimension"] == 128
    status = client.get(f"/api/projects/{project_id}").json()
    assert status["stats"]["knowledge_base"] is True


def test_local_evidence_files_are_packaged_and_require_auth():
    detail = client.get("/api/documents/nist_fips203")
    assert detail.status_code == 200
    assert detail.json()["local_original"] == "originals/NIST-FIPS-203.pdf"
    pdf = client.get("/api/evidence/nist_fips203/original")
    assert pdf.status_code == 200
    assert pdf.headers["content-type"] == "application/pdf"
    assert pdf.content.startswith(b"%PDF-")
    excerpt = client.get("/api/evidence/nist_fips203/excerpt")
    assert excerpt.status_code == 200
    assert "ML-KEM-768" in excerpt.text
    announcement = client.get("/api/evidence/gmit_2023_revision/original")
    assert announcement.status_code == 200
    assert announcement.headers["content-type"] == "text/plain; charset=utf-8"
    assert "GM/T 0009-2023" in announcement.text
    domestic_excerpt = client.get("/api/evidence/kunpeng_secgear_dev/excerpt")
    assert domestic_excerpt.status_code == 200
    assert "-DENCLAVE=GP" in domestic_excerpt.text
    assert client.get("/api/evidence/aes/original").status_code == 404
    client.post("/api/auth/logout")
    assert client.get("/api/evidence/nist_fips203/original").status_code == 401


def test_health_reports_database_and_test_reset_is_isolated():
    health = client.get("/api/health")
    assert health.status_code == 200
    assert health.json()["database"] == "ok"
    assert platform_service.repository.url.startswith("sqlite")


def test_staged_kb_pipeline_matches_final_build():
    project = client.post("/api/projects", json={"name": "分步建库"}).json()
    payload = {
        "project_id": project["id"],
        "document_ids": ["aes", "rsa"],
        "chunk_size": 256,
        "overlap": 32,
    }

    parsed = client.post(
        "/api/kb/parse",
        json={"project_id": project["id"], "document_ids": ["aes", "rsa"]},
    )
    assert parsed.status_code == 200
    assert parsed.json()["document_count"] == 2
    assert parsed.json()["documents"][0]["chars"] > 0
    assert parsed.json()["documents"][0]["preview"]

    chunked = client.post("/api/kb/chunk", json=payload)
    assert chunked.status_code == 200
    chunk_body = chunked.json()
    assert chunk_body["chunk_count"] >= 2
    assert chunk_body["max_chars"] >= chunk_body["min_chars"]

    embedded = client.post("/api/kb/embed", json=payload)
    assert embedded.status_code == 200
    embed_body = embedded.json()
    assert embed_body["dimension"] == 128
    assert len(embed_body["chunks"][0]["embedding_preview"]) == 8
    assert "离线" not in embed_body["model"]

    built = client.post("/api/kb/build", json=payload)
    assert built.status_code == 200
    assert (
        built.json()["chunk_count"]
        == chunk_body["chunk_count"]
        == embed_body["chunk_count"]
    )


def test_staged_kb_endpoints_reject_empty_selection():
    project = client.post("/api/projects", json={"name": "空选择校验"}).json()
    assert (
        client.post(
            "/api/kb/parse", json={"project_id": project["id"], "document_ids": []}
        ).status_code
        == 422
    )
    assert (
        client.post(
            "/api/kb/chunk",
            json={
                "project_id": project["id"],
                "document_ids": ["aes"],
                "chunk_size": 128,
                "overlap": 128,
            },
        ).status_code
        == 422
    )


def test_staged_rag_pipeline_stages_are_real():
    project_id, _ = create_project_and_kb()
    query = "RSA 为什么不适合直接加密 1GB 大文件？"

    embed = client.post(
        "/api/rag/embed-query", json={"project_id": project_id, "query": query}
    )
    assert embed.status_code == 200
    assert embed.json()["dimension"] == 128
    assert len(embed.json()["preview"]) == 8

    search = client.post(
        "/api/kb/search",
        json={"project_id": project_id, "query": query, "top_k": 5, "threshold": 0.02},
    )
    retrieved = search.json()["results"]
    assert retrieved

    rerank = client.post(
        "/api/rag/rerank",
        json={
            "project_id": project_id,
            "query": query,
            "items": retrieved,
            "rerank_enabled": True,
            "rerank_top_n": 3,
        },
    )
    assert rerank.status_code == 200
    assert len(rerank.json()["items"]) <= 3
    assert rerank.json()["before"] == [row["id"] for row in retrieved]

    context = client.post(
        "/api/rag/context",
        json={
            "project_id": project_id,
            "query": query,
            "items": rerank.json()["items"],
            "max_context_tokens": 1600,
        },
    )
    assert context.status_code == 200
    ctx = context.json()
    assert ctx["tokens"] <= ctx["max_tokens"]
    assert len(ctx["items"]) >= 1


def test_rsa_retrieval_and_rag_trace():
    project_id, _ = create_project_and_kb()
    query = "RSA 为什么不适合直接加密 1GB 大文件？"
    search = client.post(
        "/api/kb/search",
        json={"project_id": project_id, "query": query, "top_k": 5, "threshold": 0.02},
    )
    assert search.status_code == 200
    assert any("RSA" in row["document_title"] for row in search.json()["results"][:3])
    rag = client.post(
        "/api/rag/compare",
        json={
            "project_id": project_id,
            "query": query,
            "top_k": 5,
            "threshold": 0.02,
            "rerank_enabled": True,
            "rerank_top_n": 3,
        },
    )
    body = rag.json()
    assert rag.status_code == 200
    assert "混合加密" in body["rag"]["answer"]
    assert body["trace"]["context"]["items"]
    assert body["rag"]["citations"]


def test_private_evidence_benchmark_makes_rag_gain_measurable():
    project = client.post("/api/projects", json={"name": "RAG 证据挑战"}).json()
    built = client.post(
        "/api/kb/build",
        json={
            "project_id": project["id"],
            "document_ids": ["hailan_crypto_manual"],
            "chunk_size": 512,
            "overlap": 64,
        },
    )
    assert built.status_code == 200
    benchmark = next(
        row
        for row in client.get("/api/bootstrap").json()["rag_benchmarks"]
        if row["id"] == "private_incident"
    )
    response = client.post(
        "/api/rag/compare",
        json={
            "project_id": project["id"],
            "query": benchmark["question"],
            "benchmark_id": benchmark["id"],
            "top_k": 5,
            "threshold": 0.0,
            "rerank_enabled": True,
            "rerank_top_n": 3,
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["benchmark"]["base_score"] == 0
    assert body["benchmark"]["rag_score"] == 100
    assert body["benchmark"]["knowledge_gain"] == 100
    assert body["benchmark"]["missing_document_ids"] == []
    assert all(fact["rag_hit"] for fact in body["benchmark"]["facts"])
    assert not any(fact["base_hit"] for fact in body["benchmark"]["facts"])
    citation = body["rag"]["citations"][0]
    assert citation["source_type"] == "虚构教学情境"
    assert "教学虚构" in citation["scenario_notice"]
    assert "知识增益明显" in body["diagnosis"]


def test_standard_benchmark_exposes_verifiable_primary_source():
    project = client.post("/api/projects", json={"name": "FIPS 参数核验"}).json()
    client.post(
        "/api/kb/build",
        json={
            "project_id": project["id"],
            "document_ids": ["nist_fips203"],
            "chunk_size": 512,
            "overlap": 64,
        },
    )
    benchmark = next(
        row
        for row in client.get("/api/bootstrap").json()["rag_benchmarks"]
        if row["id"] == "fips203_sizes"
    )
    response = client.post(
        "/api/rag/compare",
        json={
            "project_id": project["id"],
            "query": benchmark["question"],
            "benchmark_id": benchmark["id"],
            "top_k": 5,
            "threshold": 0.0,
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["benchmark"]["rag_score"] == 100
    assert (
        body["rag"]["citations"][0]["source_url"]
        == "https://doi.org/10.6028/NIST.FIPS.203"
    )
    assert body["rag"]["citations"][0]["source_date"] == "2024-08-13"


def test_2026_signature_benchmark_is_grounded_in_current_nist_report():
    project = client.post("/api/projects", json={"name": "2026 签名候选核验"}).json()
    client.post(
        "/api/kb/build",
        json={
            "project_id": project["id"],
            "document_ids": ["nist_ir8610_2026"],
            "chunk_size": 512,
            "overlap": 64,
        },
    )
    benchmark = next(
        row
        for row in client.get("/api/bootstrap").json()["rag_benchmarks"]
        if row["id"] == "nist_signatures_2026"
    )
    response = client.post(
        "/api/rag/compare",
        json={
            "project_id": project["id"],
            "query": benchmark["question"],
            "benchmark_id": benchmark["id"],
            "top_k": 5,
            "threshold": 0.0,
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["benchmark"]["base_score"] == 0
    assert body["benchmark"]["rag_score"] == 100
    assert (
        body["rag"]["citations"][0]["source_url"]
        == "https://doi.org/10.6028/NIST.IR.8610"
    )


@pytest.mark.parametrize(
    ("benchmark_id", "document_ids"),
    [
        ("gmt_2023_revision", ["gmit_2023_revision"]),
        ("gmit_eval_stack", ["cn_crypto_baseline_gbt39786", "gmit_eval_2021"]),
        ("kunpeng_secgear_build", ["kunpeng_secgear_dev"]),
        (
            "phytium_security_stack",
            ["phytium_phytee_platform", "phytium_tee_architecture"],
        ),
    ],
)
def test_domestic_standard_and_tee_benchmarks_show_grounded_gain(
    benchmark_id, document_ids
):
    project = client.post(
        "/api/projects", json={"name": f"国产专题核验 {benchmark_id}"}
    ).json()
    built = client.post(
        "/api/kb/build",
        json={
            "project_id": project["id"],
            "document_ids": document_ids,
            "chunk_size": 512,
            "overlap": 64,
        },
    )
    assert built.status_code == 200
    benchmark = next(
        row
        for row in client.get("/api/bootstrap").json()["rag_benchmarks"]
        if row["id"] == benchmark_id
    )
    response = client.post(
        "/api/rag/compare",
        json={
            "project_id": project["id"],
            "query": benchmark["question"],
            "benchmark_id": benchmark_id,
            "top_k": 8,
            "threshold": -0.1,
            "rerank_enabled": True,
            "rerank_top_n": 6,
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["benchmark"]["base_score"] == 0
    assert body["benchmark"]["rag_score"] == 100
    assert body["benchmark"]["knowledge_gain"] == 100
    assert body["benchmark"]["missing_document_ids"] == []
    assert all(fact["rag_hit"] for fact in body["benchmark"]["facts"])


def test_internal_chunk_markers_are_removed_from_student_answer():
    answer = '第一条结论。（来源：`chk_nist_ir8610_2026_002`）\n第二条（chunk_id="chk_demo_003"）\n根据 `chk_demo_004` 可知。'
    cleaned = platform_service._strip_internal_source_markers(answer)
    assert cleaned == "第一条结论。\n第二条\n根据 知识库证据 可知。"
    assert "chk_" not in cleaned


def test_benchmark_reports_missing_required_document_and_rejects_edited_question():
    project_id, _ = create_project_and_kb()
    benchmark = next(
        row
        for row in client.get("/api/bootstrap").json()["rag_benchmarks"]
        if row["id"] == "private_incident"
    )
    missing = client.post(
        "/api/rag/compare",
        json={
            "project_id": project_id,
            "query": benchmark["question"],
            "benchmark_id": benchmark["id"],
            "top_k": 3,
            "threshold": 0.0,
        },
    )
    assert missing.status_code == 200
    assert missing.json()["benchmark"]["missing_document_ids"] == [
        "hailan_crypto_manual"
    ]
    edited = client.post(
        "/api/rag/compare",
        json={
            "project_id": project_id,
            "query": benchmark["question"] + "（已改）",
            "benchmark_id": benchmark["id"],
        },
    )
    assert edited.status_code == 400


def test_agent_selection_trace_and_safe_tools():
    project_id, _ = create_project_and_kb()
    response = client.post(
        "/api/agents/run",
        json={
            "project_id": project_id,
            "query": "我需要在云服务器处理敏感数据，请比较 HE、MPC、TEE 并给出选型建议。",
            "memory_enabled": True,
        },
    )
    body = response.json()
    assert response.status_code == 200
    assert body["selected_skill"]["id"] == "crypto_selection"
    assert body["tool_calls"][0]["tool"] == "knowledge_search"
    assert body["metrics"]["status"] == "completed"
    calculator = client.post(
        "/api/tools/calculator/run", json={"arguments": {"expression": "20 / 1.25"}}
    )
    assert calculator.json()["output"]["value"] == 16
    blocked = client.post(
        "/api/tools/calculator/run",
        json={"arguments": {"expression": "__import__('os').system('id')"}},
    )
    assert blocked.status_code == 400


def test_multiple_projects_can_each_build_same_agent():
    first_id, _ = create_project_and_kb()
    second_id, _ = create_project_and_kb()
    payload = {"query": "比较 HE、MPC、TEE 并给出选型建议", "memory_enabled": True}
    assert (
        client.patch(f"/api/projects/{first_id}", json={"ended": False}).status_code
        == 200
    )
    first = client.post("/api/agents/run", json={"project_id": first_id, **payload})
    assert (
        client.patch(f"/api/projects/{second_id}", json={"ended": False}).status_code
        == 200
    )
    second = client.post("/api/agents/run", json={"project_id": second_id, **payload})
    assert first.status_code == 200
    assert second.status_code == 200
    assert client.get(f"/api/projects/{first_id}").json()["stats"]["agent"] is True
    assert client.get(f"/api/projects/{second_id}").json()["stats"]["agent"] is True


def test_report_after_complete_flow():
    project_id, _ = create_project_and_kb()
    client.post(
        "/api/rag/compare",
        json={
            "project_id": project_id,
            "query": "TrustZone 与 SGX 的隔离机制有什么区别？",
            "top_k": 5,
            "threshold": 0.0,
        },
    )
    client.post(
        "/api/agents/run",
        json={
            "project_id": project_id,
            "query": "比较 HE、MPC、TEE 并给出选型建议",
            "memory_enabled": True,
        },
    )
    report = client.get(f"/api/reports/{project_id}")
    assert report.status_code == 200
    assert len(report.json()["runs"]) >= 3
    assert "能力画像" in report.json()["markdown"]


def test_report_exports_real_pdf_and_docx_with_conclusion():
    project_id, _ = create_project_and_kb()
    conclusion = "调整 Top-K 后，引用覆盖率提高，但也需要控制上下文噪声。"

    pdf = client.post(
        f"/api/reports/{project_id}/export",
        json={"format": "pdf", "conclusion": conclusion},
    )
    assert pdf.status_code == 200
    assert pdf.headers["content-type"] == "application/pdf"
    assert ".pdf" in pdf.headers["content-disposition"]
    assert pdf.content.startswith(b"%PDF-")
    assert len(PdfReader(BytesIO(pdf.content)).pages) >= 1

    docx = client.post(
        f"/api/reports/{project_id}/export",
        json={"format": "docx", "conclusion": conclusion},
    )
    assert docx.status_code == 200
    assert (
        docx.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert ".docx" in docx.headers["content-disposition"]
    assert docx.content.startswith(b"PK")
    document = Document(BytesIO(docx.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "验收项目 · 密码学实验报告" in text
    assert conclusion in text


def test_rich_text_observation_is_sanitized_persisted_and_exported_with_formatting():
    project_id, _ = create_project_and_kb()
    observation_html = (
        "<h3>关键观察</h3><p><strong>Top-K</strong> 提升后，<em>引用覆盖率</em>提高。</p>"
        "<blockquote>RAG 增益来自可核验证据。</blockquote>"
        "<ul><li>继续控制上下文噪声</li></ul>"
        '<p><a href="https://example.com/evidence">实验参考</a>'
        "<a href=\"javascript:alert(1)\">危险链接</a></p><script>alert('x')</script>"
    )
    saved = client.put(
        f"/api/reports/{project_id}/observation", json={"html": observation_html}
    )
    assert saved.status_code == 200
    body = saved.json()
    assert "<script" not in body["observation"]["html"]
    assert "javascript:" not in body["observation"]["html"]
    assert "<strong>Top-K</strong>" in body["observation"]["html"]
    assert body["observation"]["updated_at"]
    assert "### 关键观察" in body["observation"]["markdown"]
    assert "**Top-K**" in body["markdown"]
    assert "## 观察和感想" in body["markdown"]

    persisted = platform_service.repository.load()["projects"][project_id]
    assert persisted["observation_html"] == body["observation"]["html"]

    pdf = client.post(f"/api/reports/{project_id}/export", json={"format": "pdf"})
    assert pdf.status_code == 200
    assert pdf.content.startswith(b"%PDF-")
    pdf_text = "\n".join(
        page.extract_text() or "" for page in PdfReader(BytesIO(pdf.content)).pages
    )
    assert "观察和感想" in pdf_text
    assert "Top-K" in pdf_text

    docx = client.post(f"/api/reports/{project_id}/export", json={"format": "docx"})
    assert docx.status_code == 200
    document = Document(BytesIO(docx.content))
    paragraphs = document.paragraphs
    assert any("七、观察和感想" in paragraph.text for paragraph in paragraphs)
    assert any(
        run.bold and "Top-K" in run.text
        for paragraph in paragraphs
        for run in paragraph.runs
    )
    assert any(
        paragraph.style.name == "List Bullet" and "继续控制上下文噪声" in paragraph.text
        for paragraph in paragraphs
    )


def test_report_export_rejects_unknown_format():
    project = client.post("/api/projects", json={"name": "格式校验"}).json()
    response = client.post(
        f"/api/reports/{project['id']}/export",
        json={"format": "doc", "conclusion": ""},
    )
    assert response.status_code == 422


def test_upload_text_and_code_then_build_knowledge_base():
    project = client.post("/api/projects", json={"name": "用户资料实验"}).json()
    text_upload = client.post(
        "/api/documents/upload",
        data={"project_id": project["id"]},
        files={
            "file": (
                "my-notes.txt",
                "格密码的安全性通常依赖格上困难问题。",
                "text/plain",
            )
        },
    )
    assert text_upload.status_code == 201
    assert text_upload.json()["source"] == "upload"

    code_upload = client.post(
        "/api/documents/upload",
        data={"project_id": project["id"]},
        files={
            "file": (
                "rsa_demo.py",
                b"def mod_pow(a, e, n):\n    return pow(a, e, n)\n",
                "text/x-python",
            )
        },
    )
    assert code_upload.status_code == 201
    code = code_upload.json()
    assert code["file_kind"] == "code"
    assert code["language"] == "Python"

    detail = client.get(f"/api/documents/{code['id']}?project_id={project['id']}")
    assert detail.status_code == 200
    assert detail.json()["parsed"]["language"] == "Python"
    assert "storage_path" not in detail.json()

    built = client.post(
        "/api/kb/build",
        json={
            "project_id": project["id"],
            "document_ids": [text_upload.json()["id"], code["id"]],
            "chunk_size": 128,
            "overlap": 0,
        },
    )
    assert built.status_code == 200
    assert built.json()["document_count"] == 2
    assert any(
        "代码符号：mod_pow" == chunk["section"] for chunk in built.json()["chunks"]
    )


def test_delete_uploaded_document_removes_it_and_rejects_presets():
    project = client.post("/api/projects", json={"name": "删除资料实验"}).json()
    upload = client.post(
        "/api/documents/upload",
        data={"project_id": project["id"]},
        files={"file": ("temp-notes.txt", "临时上传的密码学笔记。", "text/plain")},
    )
    assert upload.status_code == 201
    document_id = upload.json()["id"]
    assert any(
        d["id"] == document_id
        for d in client.get(f"/api/projects/{project['id']}/documents").json()
    )

    deleted = client.request(
        "DELETE", f"/api/documents/{document_id}", json={"project_id": project["id"]}
    )
    assert deleted.status_code == 200
    assert deleted.json()["deleted"] == document_id
    assert not any(
        d["id"] == document_id
        for d in client.get(f"/api/projects/{project['id']}/documents").json()
    )

    preset = client.request(
        "DELETE", "/api/documents/aes", json={"project_id": project["id"]}
    )
    assert preset.status_code in (400, 404)


def test_upload_rejects_unsupported_binary_and_oversize():
    project = client.post("/api/projects", json={"name": "上传安全测试"}).json()
    unsupported = client.post(
        "/api/documents/upload",
        data={"project_id": project["id"]},
        files={"file": ("attack.exe", b"MZ", "application/octet-stream")},
    )
    assert unsupported.status_code == 400

    fake_pdf = client.post(
        "/api/documents/upload",
        data={"project_id": project["id"]},
        files={"file": ("fake.pdf", b"not a pdf", "application/pdf")},
    )
    assert fake_pdf.status_code == 400

    oversized = client.post(
        "/api/documents/upload",
        data={"project_id": project["id"]},
        files={"file": ("large.txt", b"a" * (10 * 1024 * 1024 + 1), "text/plain")},
    )
    assert oversized.status_code == 400


def test_upload_pdf_extracts_page_text():
    project = client.post("/api/projects", json={"name": "PDF 解析实验"}).json()
    buffer = BytesIO()
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
    )
    stream = StreamObject()
    stream.set_data(
        b"BT /F1 12 Tf 72 720 Td (RSA hybrid encryption uses AES for bulk data.) Tj ET"
    )
    page[NameObject("/Contents")] = writer._add_object(stream)
    writer.write(buffer)

    response = client.post(
        "/api/documents/upload",
        data={"project_id": project["id"]},
        files={"file": ("crypto-paper.pdf", buffer.getvalue(), "application/pdf")},
    )
    assert response.status_code == 201
    body = response.json()
    assert body["file_kind"] == "pdf"
    assert body["pages"] == 1
    detail = client.get(
        f"/api/documents/{body['id']}?project_id={project['id']}"
    ).json()
    assert "hybrid encryption" in detail["content"]
    assert detail["parsed"]["pages"] == 1
